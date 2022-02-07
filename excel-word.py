from openpyxl import load_workbook, Workbook
from openpyxl.styles import PatternFill, Font
from os.path import join, abspath
from docx import Document
from docx.shared import Cm, Inches
from docx.shared import Pt


def extract_excel():
	print('Введите название файла')
	filename = input()

	path = "source_files/" + filename

	data_path = join('.', path)
	data_path = abspath(data_path)

	wb = load_workbook(path, data_only=True, read_only=True)

	tables_name = list(wb.sheetnames)

	list_1 = wb[tables_name[0]]

	name_row = [cell.value for cell in next(
		list_1.iter_rows(min_row=1, min_col=1, max_row=1, max_col=list_1.max_column))]

	to_print = list()
	for i in range(2, list_1.max_row + 1):
		to_print.append([cell.value for cell in next(
		list_1.iter_rows(min_row=i, min_col=1, max_row=i, max_col=list_1.max_column))])

	return to_print, filename



def make_first_row_in_cell_bold(table):
	for i in range(1):
		for j in range(len(table.columns)):
        		cell = table.cell(i, j)
		        for paragraph in cell.paragraphs:
                		for run in paragraph.runs:
                                      run.font.bold = True
                                      break
			


def extract_excel_file_to_word_tables(extract_values, filename):
	document = Document()

	print(len(extract_values))

	for i in range(len(extract_values)-7):
		values = extract_values[i]
		
		table = document.add_table(rows=5, cols=4, style='Table Grid')
		p = document.add_paragraph("")
		table.autofit = False


		# Set properties of cells
		list_properties = [0.5, 0.5, 0.9, 0.7]
		for i in range(0, 4):
			for cell in table.rows[0].cells:
				cell.width = Inches(list_properties[i])


		# Set structure of table
		First = table.cell(0,0).merge(table.cell(1,0))
		Second= table.cell(0,1).merge(table.cell(0,2))
		Third = table.cell(1,1).merge(table.cell(1,2))
		Fourth = table.cell(3,0).merge(table.cell(3,1))
		Fieth= table.cell(3,1).merge(table.cell(3,2))
		Sixth = table.cell(3,2).merge(table.cell(3,3))


		# Set cells name
		First.text = str(values[0])

		Second.text = "Название"
		table.cell(1, 1).text = values[2]

		table.cell(0, 3).add_paragraph("Тип проблемы").style.font.bold = True
		par = table.cell(2, 0).paragraphs[0]
		par.style.font.bold = True
		table.cell(2,0).text = "Создатель:"
		par_1 = table.cell(2, 0).add_paragraph(values[6])
		par_1.style.font.bold = False
		
		table.cell(2, 1).text = "Дата создания:\n" + str(values[13])
		table.cell(2, 2).text = "Исполнитель:\n" + values[8]
		table.cell(2, 3).text = "Ответственный\n"
		table.cell(3, 0).text = "Краткое описание:\n" + values[4]
		table.cell(4, 0).text = "Тестировщик:\n"
		table.cell(4, 1).text = "Дата проверки:\n"
		version = str(values[15])
		table.cell(4, 2).text = "Протестировано в версии:\n" + version[2:]
		table.cell(4, 3).text = "Дата закрытия:\n" + str(values[14])


		# Set properties for each cell
		make_first_row_in_cell_bold(table)
		document.save("extraction_result/" + filename + ".docx")

	document.save("extraction_result/" + filename + ".docx")


n, file = extract_excel()
extract_excel_file_to_word_tables(n, file)
